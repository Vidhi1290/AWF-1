"""
Generates the three sample resume files (DOCX/PDF) used for testing the
workflow. Run once: python3 generate_samples.py

Requires: python-docx (pip install python-docx)
For the PDF sample, also requires LibreOffice ("soffice") on PATH to convert
a DOCX to a standards-compliant PDF (most real-world resume PDFs are produced
the same way - exported from Word/Google Docs - so this keeps the sample
realistic). If "soffice" isn't available, the script falls back to writing
the PDF with reportlab, but note some minimal-xref PDFs written by simpler
generators have been observed to trip up older PDF parsers (e.g. pdf-parse
1.1.1) - if you hit "bad XRef entry" during testing, regenerate the PDF via
Word/Google Docs/LibreOffice export instead.
"""
import os
import shutil
import subprocess
from docx import Document

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

RESUMES = {
    "resume_strong_match_priya_sharma": {
        "format": "docx",
        "lines": [
            ("title", "Priya Sharma"),
            ("sub", "priya.sharma.dev@gmail.com | +91 98765 43210 | Bengaluru, India"),
            ("sub", "linkedin.com/in/priyasharma-dev"),
            ("h", "Summary"),
            ("p", "Backend engineer with 5 years of experience building scalable Node.js and Python "
                  "services, REST APIs, and automation pipelines. Strong background in cloud "
                  "infrastructure and CI/CD."),
            ("h", "Experience"),
            ("p", "Senior Backend Engineer — Fintrix Labs (2022 - Present)"),
            ("p", "- Designed and built REST APIs in Node.js serving 2M+ requests/day."),
            ("p", "- Migrated legacy cron jobs to n8n workflow automation, reducing manual ops by 70%."),
            ("p", "- Integrated OpenAI and open-source LLMs for document processing pipelines."),
            ("p", "- Managed Postgres database design and query optimization."),
            ("p", "Backend Engineer — Codebase Solutions (2020 - 2022)"),
            ("p", "- Built microservices in Python (FastAPI) deployed via Docker on AWS ECS."),
            ("p", "- Set up GitHub Actions CI/CD pipelines for automated testing and deployment."),
            ("p", "- Implemented Kafka-based event pipeline for order processing."),
            ("h", "Education"),
            ("p", "B.Tech in Computer Science, VIT Vellore (2016 - 2020)"),
            ("h", "Technical Skills"),
            ("p", "Node.js, Python, FastAPI, Express, PostgreSQL, MySQL, Docker, AWS, GCP, "
                  "GitHub Actions, Kafka, n8n, REST API design, Git"),
            ("h", "Soft Skills"),
            ("p", "Communication, Ownership, Mentoring, Cross-team collaboration"),
            ("h", "Certifications"),
            ("p", "AWS Certified Solutions Architect - Associate"),
            ("h", "Languages"),
            ("p", "English (Fluent), Hindi (Native), Kannada (Conversational)"),
            ("h", "Projects"),
            ("p", "Resume Screening Bot - internal tool using LLMs to pre-screen job applications."),
            ("p", "Order Pipeline Revamp - Kafka-based event-driven order processing system."),
        ],
    },
    "resume_medium_match_arjun_verma": {
        "format": "docx",
        "lines": [
            ("title", "Arjun Verma"),
            ("sub", "arjun.verma91@gmail.com | 9123456780 | Pune, India"),
            ("h", "Summary"),
            ("p", "Software developer with 3 years of experience, primarily in Python scripting and "
                  "data processing. Some exposure to backend APIs and cloud basics."),
            ("h", "Experience"),
            ("p", "Software Developer — DataWorks Analytics (2021 - Present)"),
            ("p", "- Built internal Python scripts for data cleaning and ETL."),
            ("p", "- Wrote Flask endpoints for internal dashboards."),
            ("p", "- Used MySQL for reporting queries."),
            ("p", "Junior Developer — StartHub Technologies (2020 - 2021)"),
            ("p", "- Assisted in building a Django-based internal tool."),
            ("p", "- Basic exposure to Git and version control."),
            ("h", "Education"),
            ("p", "B.Sc in Information Technology, Pune University (2017 - 2020)"),
            ("h", "Technical Skills"),
            ("p", "Python, Flask, Django, MySQL, Git, Pandas, basic AWS (S3, EC2)"),
            ("h", "Soft Skills"),
            ("p", "Teamwork, Adaptability, Problem-solving"),
            ("h", "Certifications"),
            ("p", "None listed"),
            ("h", "Languages"),
            ("p", "English (Fluent), Hindi (Native), Marathi (Native)"),
            ("h", "Projects"),
            ("p", "Sales Dashboard - internal reporting tool built with Flask and MySQL."),
        ],
    },
    "resume_weak_match_neha_kapoor": {
        "format": "pdf",
        "lines": [
            ("title", "Neha Kapoor"),
            ("sub", "neha.kapoor.design@gmail.com | +91 90000 11223 | Mumbai, India"),
            ("h", "Summary"),
            ("p", "Graphic designer with 4 years of experience in branding, UI mockups, and marketing "
                  "collateral. Recently exploring no-code tools out of personal interest."),
            ("h", "Experience"),
            ("p", "Senior Graphic Designer - Canvas Creative Studio (2021 - Present)"),
            ("p", "- Designed branding assets and marketing materials for clients."),
            ("p", "- Created UI mockups in Figma for client mobile apps."),
            ("p", "- Coordinated with print vendors for packaging design."),
            ("p", "Junior Designer - PixelWorks (2020 - 2021)"),
            ("p", "- Produced social media graphics and ad creatives."),
            ("h", "Education"),
            ("p", "B.Des in Visual Communication, NIFT Mumbai (2016 - 2020)"),
            ("h", "Technical Skills"),
            ("p", "Figma, Adobe Photoshop, Adobe Illustrator, Canva, basic HTML/CSS"),
            ("h", "Soft Skills"),
            ("p", "Creativity, Attention to detail, Client communication"),
            ("h", "Certifications"),
            ("p", "Google UX Design Certificate"),
            ("h", "Languages"),
            ("p", "English (Fluent), Hindi (Native)"),
            ("h", "Projects"),
            ("p", "Rebrand of a D2C skincare startup - full brand identity package."),
        ],
    },
}


def build_docx(path, lines):
    doc = Document()
    for kind, text in lines:
        if kind == "title":
            doc.add_heading(text, level=0)
        elif kind == "sub":
            doc.add_paragraph(text)
        elif kind == "h":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(path)


def convert_docx_to_pdf_with_soffice(docx_path, out_dir):
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        check=True, capture_output=True,
    )


def build_pdf_with_reportlab(path, lines):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    story = []
    for kind, text in lines:
        story.append(Paragraph(text, styles["Title"] if kind == "title" else
                                (styles["Heading2"] if kind == "h" else styles["Normal"])))
        story.append(Spacer(1, 4))
    doc.build(story)


if __name__ == "__main__":
    for name, spec in RESUMES.items():
        if spec["format"] == "docx":
            build_docx(os.path.join(OUT_DIR, f"{name}.docx"), spec["lines"])
            print("wrote", f"{name}.docx")
        else:
            tmp_docx = os.path.join(OUT_DIR, f"_tmp_{name}.docx")
            build_docx(tmp_docx, spec["lines"])
            if shutil.which("soffice"):
                convert_docx_to_pdf_with_soffice(tmp_docx, OUT_DIR)
                generated_pdf = os.path.join(OUT_DIR, f"_tmp_{name}.pdf")
                target_pdf = os.path.join(OUT_DIR, f"{name}.pdf")
                shutil.move(generated_pdf, target_pdf)
                print("wrote", f"{name}.pdf", "(via LibreOffice)")
            else:
                target_pdf = os.path.join(OUT_DIR, f"{name}.pdf")
                build_pdf_with_reportlab(target_pdf, spec["lines"])
                print("wrote", f"{name}.pdf", "(via reportlab fallback - soffice not found)")
            try:
                os.remove(tmp_docx)
            except OSError:
                pass  # best-effort cleanup; harmless if it lingers
